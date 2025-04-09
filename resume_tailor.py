import os
import sys
import json
import argparse
from typing import Dict, List, Tuple, Any, Optional, Literal
import re
import textwrap
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# LangChain imports
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.output_parsers import PydanticOutputParser
from langchain.pydantic_v1 import BaseModel, Field
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

# LLM providers
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_anthropic import ChatAnthropic

class JobAnalysis(BaseModel):
    """Analysis of a job description for resume tailoring."""
    keywords: List[str] = Field(description="The most important technical skills, tools, and technologies mentioned")
    soft_skills: List[str] = Field(description="Soft skills and qualities emphasized")
    domain_knowledge: List[str] = Field(description="Specific industry or domain knowledge required")
    experience_focus: List[str] = Field(description="Aspects of work experience that should be emphasized")
    tech_stack: List[str] = Field(description="Primary technologies that form the core tech stack")
    related_technologies: List[str] = Field(description="Technologies not explicitly mentioned but likely relevant")
    priority_order: List[str] = Field(description="Suggested order of importance for skills to highlight")
    tone: str = Field(description="Suggested tone for the resume (e.g., technical, business-focused, innovative)")

class ResumeTailor:
    def __init__(self, cv_dir: str = "cv", llm_provider: str = "openai", api_key: Optional[str] = None):
        """
        Initialize the ResumeTailor with the directory containing resume components.
        
        Args:
            cv_dir: Directory containing resume component files
            llm_provider: LLM provider to use ('openai', 'google', or 'anthropic')
            api_key: API key for the selected LLM provider
        """
        self.cv_dir = cv_dir
        self.resume_components = {
            "professional_summary": self._read_file("ps.txt"),
            "skills": self._read_file("skills.txt"),
            "experience": self._read_file("experience.txt"),
            "education": self._read_file("education.txt")
        }
        
        # Initialize LLM based on provider
        self.llm_provider = llm_provider
        self.llm = self._initialize_llm(llm_provider, api_key)

    def _read_file(self, filename: str) -> str:
        """Read a file from the CV directory."""
        filepath = os.path.join(self.cv_dir, filename)
        try:
            with open(filepath, 'r') as file:
                return file.read()
        except FileNotFoundError:
            print(f"Warning: File {filepath} not found.")
            return ""

    def _initialize_llm(self, provider: str, api_key: Optional[str] = None):
        """Initialize the LLM based on the provider."""
        if provider == "openai":
            return ChatOpenAI(
                model="gpt-4",
                temperature=0.2,
                api_key=api_key or os.environ.get("OPENAI_API_KEY")
            )
        elif provider == "google":
            return ChatGoogleGenerativeAI(
                model="gemini-pro",
                temperature=0.2,
                api_key=api_key or os.environ.get("GOOGLE_API_KEY")
            )
        elif provider == "anthropic":
            return ChatAnthropic(
                model="claude-3-opus-20240229",
                temperature=0.2,
                api_key=api_key or os.environ.get("ANTHROPIC_API_KEY")
            )
        else:
            raise ValueError(f"Unsupported LLM provider: {provider}")

    def extract_keywords(self, job_description: str) -> Dict[str, Any]:
        """
        Extract keywords from the job description using LangChain and LLM.
        Returns a dictionary with categorized keywords and other analysis.
        """
        # Create a parser for structured output
        parser = PydanticOutputParser(pydantic_object=JobAnalysis)
        
        # Create a prompt template
        prompt_template = PromptTemplate(
            template="""
            You are a skilled recruiter and resume expert.
            
            Analyze the following job description and extract key information:
            
            {job_description}
            
            {format_instructions}
            """,
            input_variables=["job_description"],
            partial_variables={"format_instructions": parser.get_format_instructions()}
        )
        
        # Create the chain
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        try:
            # Run the chain
            result = chain.run(job_description=job_description)
            
            # Parse the result
            parsed_result = parser.parse(result)
            
            # Convert to dictionary
            return parsed_result.dict()
            
        except Exception as e:
            print(f"Error extracting keywords: {e}")
            # Return a basic structure if the chain fails
            return {
                "keywords": [],
                "soft_skills": [],
                "domain_knowledge": [],
                "experience_focus": [],
                "tech_stack": [],
                "related_technologies": [],
                "priority_order": [],
                "tone": "balanced"
            }

    def tailor_professional_summary(self, analysis: Dict[str, Any]) -> str:
        """Tailor the professional summary based on job analysis using LangChain."""
        prompt_template = PromptTemplate(
            template="""
            You are a skilled resume writer.
            
            Tailor this professional summary for a job that emphasizes: 
            - Keywords: {keywords}
            - Tech stack: {tech_stack}
            - Domain knowledge: {domain_knowledge}
            - Tone: {tone}
            
            Original summary:
            {original_summary}
            
            Provide a tailored professional summary that:
            1. Maintains the same length as the original
            2. Emphasizes relevant experience and skills for this specific job
            3. Uses a tone appropriate for the role
            4. Includes key technologies from the job description where relevant
            5. Avoids generic statements and focuses on specific value
            
            Return only the revised summary text.
            """,
            input_variables=["keywords", "tech_stack", "domain_knowledge", "tone", "original_summary"]
        )
        
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        try:
            result = chain.run(
                keywords=', '.join(analysis.get('keywords', [])),
                tech_stack=', '.join(analysis.get('tech_stack', [])),
                domain_knowledge=', '.join(analysis.get('domain_knowledge', [])),
                tone=analysis.get('tone', 'balanced'),
                original_summary=self.resume_components['professional_summary']
            )
            
            return result.strip()
            
        except Exception as e:
            print(f"Error tailoring professional summary: {e}")
            return self.resume_components['professional_summary']

    def tailor_skills(self, analysis: Dict[str, Any]) -> str:
        """Tailor the skills section based on job analysis using LangChain."""
        # Create a mapping of skills to their priority based on the job analysis
        priority_skills = {skill.lower(): i for i, skill in enumerate(analysis.get('priority_order', []))}
        
        # Get all skills from the original skills section
        skills_text = self.resume_components['skills']
        
        prompt_template = PromptTemplate(
            template="""
            You are a skilled resume writer.
            
            Tailor this skills section for a job that emphasizes: 
            - Keywords: {keywords}
            - Tech stack: {tech_stack}
            - Related technologies: {related_technologies}
            - Soft skills: {soft_skills}
            
            Original skills:
            {original_skills}
            
            Provide a tailored skills section that:
            1. Prioritizes skills that match the job requirements
            2. Keeps the bullet point format (starting with •)
            3. Groups related skills together
            4. Intelligently substitutes technologies when appropriate (e.g., if the job mentions AWS but the resume has GCP experience)
            5. Emphasizes transferable skills for requirements not directly matched
            6. Limits to the most relevant skills to keep the section concise
            7. Maintains the same formatting style as the original
            
            Return only the revised skills text with bullet points.
            """,
            input_variables=["keywords", "tech_stack", "related_technologies", "soft_skills", "original_skills"]
        )
        
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        try:
            result = chain.run(
                keywords=', '.join(analysis.get('keywords', [])),
                tech_stack=', '.join(analysis.get('tech_stack', [])),
                related_technologies=', '.join(analysis.get('related_technologies', [])),
                soft_skills=', '.join(analysis.get('soft_skills', [])),
                original_skills=skills_text
            )
            
            return result.strip()
            
        except Exception as e:
            print(f"Error tailoring skills: {e}")
            return skills_text

    def tailor_experience(self, analysis: Dict[str, Any]) -> str:
        """Tailor the experience section based on job analysis using LangChain."""
        prompt_template = PromptTemplate(
            template="""
            You are a skilled resume writer.
            
            Tailor this work experience section for a job that emphasizes: 
            - Keywords: {keywords}
            - Tech stack: {tech_stack}
            - Experience focus: {experience_focus}
            - Domain knowledge: {domain_knowledge}
            
            Original experience:
            {original_experience}
            
            Provide a tailored experience section that:
            1. Emphasizes projects and responsibilities most relevant to the target job
            2. Highlights achievements that demonstrate skills mentioned in the job description
            3. Adjusts technical terminology to match the job description where appropriate
            4. Maintains the same formatting and structure as the original
            5. Prioritizes the most relevant experiences
            6. Keeps the section concise while preserving important details
            7. Intelligently substitutes technologies when appropriate (e.g., if the job mentions AWS but the resume has GCP experience)
            
            Return only the revised experience text.
            """,
            input_variables=["keywords", "tech_stack", "experience_focus", "domain_knowledge", "original_experience"]
        )
        
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        try:
            result = chain.run(
                keywords=', '.join(analysis.get('keywords', [])),
                tech_stack=', '.join(analysis.get('tech_stack', [])),
                experience_focus=', '.join(analysis.get('experience_focus', [])),
                domain_knowledge=', '.join(analysis.get('domain_knowledge', [])),
                original_experience=self.resume_components['experience']
            )
            
            return result.strip()
            
        except Exception as e:
            print(f"Error tailoring experience: {e}")
            return self.resume_components['experience']

    def generate_tailored_resume(self, job_description: str) -> str:
        """
        Generate a tailored resume based on the job description.
        Returns the complete tailored resume as a string.
        """
        # Extract keywords and analyze the job description
        analysis = self.extract_keywords(job_description)
        
        # Tailor each section of the resume
        tailored_summary = self.tailor_professional_summary(analysis)
        tailored_skills = self.tailor_skills(analysis)
        tailored_experience = self.tailor_experience(analysis)
        education = self.resume_components['education']  # Education typically stays the same
        
        # Combine the sections into a complete resume
        resume = f"""# PROFESSIONAL SUMMARY
{tailored_summary}

# SKILLS
{tailored_skills}

# WORK EXPERIENCE
{tailored_experience}

# EDUCATION
{education}
"""
        
        # Ensure the resume fits within 2 pages
        resume = self.ensure_page_limit(resume)
        
        return resume

    def ensure_page_limit(self, resume: str) -> str:
        """
        Ensure the resume fits within 2 pages using LangChain.
        This is an approximation as actual page count depends on formatting.
        """
        # Approximate characters per page (based on standard formatting)
        # This is a rough estimate - actual page count depends on font, margins, etc.
        chars_per_page = 3500
        max_chars = chars_per_page * 2
        
        if len(resume) <= max_chars:
            return resume
        
        # If resume is too long, ask the LLM to condense it
        prompt_template = PromptTemplate(
            template="""
            You are a skilled resume editor.
            
            The following resume is too long (exceeds 2 pages). Please condense it while preserving the most important information:
            
            {resume}
            
            Guidelines for condensing:
            1. Maintain all section headers (Professional Summary, Skills, Work Experience, Education)
            2. Preserve the most relevant skills and experiences for the job
            3. Remove redundant or less important details
            4. Shorten descriptions while maintaining key achievements
            5. Keep the same overall structure and formatting
            6. Ensure the final resume fits within 2 pages (approximately {max_chars} characters)
            
            Return only the condensed resume.
            """,
            input_variables=["resume", "max_chars"]
        )
        
        chain = LLMChain(llm=self.llm, prompt=prompt_template)
        
        try:
            result = chain.run(resume=resume, max_chars=max_chars)
            return result.strip()
            
        except Exception as e:
            print(f"Error condensing resume: {e}")
            # If chain fails, do a simple truncation (not ideal but better than nothing)
            return resume[:max_chars] + "\n\n[Resume truncated to fit 2 pages]"

    def save_resume(self, resume: str, output_path: str, format: str = "docx") -> None:
        """
        Save the tailored resume to a file.
        
        Args:
            resume: The resume content as a string
            output_path: The path to save the resume to
            format: The format to save the resume in ('docx', 'txt', or 'md')
        """
        if format == "docx":
            self._save_as_docx(resume, output_path)
        else:
            # Save as plain text or markdown
            with open(output_path, 'w') as file:
                file.write(resume)
            print(f"Tailored resume saved to {output_path}")
    
    def _save_as_docx(self, resume: str, output_path: str) -> None:
        """Save the resume as a Word document."""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Split the resume into sections
        sections = resume.split("\n# ")
        
        # Process the first section (which doesn't start with #)
        if sections[0].startswith("# "):
            sections[0] = sections[0][2:]  # Remove the # prefix
        
        # Process each section
        for i, section in enumerate(sections):
            if i > 0:
                # For sections after the first one, the # was removed in the split
                section_parts = section.split("\n", 1)
                section_title = section_parts[0]
                section_content = section_parts[1] if len(section_parts) > 1 else ""
            else:
                # For the first section
                section_parts = section.split("\n", 1)
                section_title = section_parts[0]
                section_content = section_parts[1] if len(section_parts) > 1 else ""
            
            # Add section title
            title = doc.add_heading(section_title, level=1)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            # Add section content
            for paragraph_text in section_content.split("\n\n"):
                if paragraph_text.strip():
                    # Check if this is a bullet point list
                    if "•" in paragraph_text:
                        for bullet_point in paragraph_text.split("•"):
                            if bullet_point.strip():
                                p = doc.add_paragraph()
                                p.style = "List Bullet"
                                p.add_run(bullet_point.strip())
                    else:
                        p = doc.add_paragraph()
                        p.add_run(paragraph_text.strip())
        
        # Save the document
        doc.save(output_path)
        print(f"Tailored resume saved as Word document to {output_path}")

def main():
    parser = argparse.ArgumentParser(description='Tailor a resume based on a job description.')
    parser.add_argument('--job', type=str, help='Path to job description file')
    parser.add_argument('--job_text', type=str, help='Job description text')
    parser.add_argument('--output', type=str, default='tailored_resume.docx', help='Output file path')
    parser.add_argument('--cv_dir', type=str, default='cv', help='Directory containing resume components')
    parser.add_argument('--llm', type=str, choices=['openai', 'google', 'anthropic'], default='openai',
                        help='LLM provider to use (default: openai)')
    parser.add_argument('--api_key', type=str, help='API key for the selected LLM provider')
    parser.add_argument('--format', type=str, choices=['docx', 'txt', 'md'], default='docx',
                        help='Output format (default: docx)')
    
    args = parser.parse_args()
    
    # Initialize the ResumeTailor
    tailor = ResumeTailor(cv_dir=args.cv_dir, llm_provider=args.llm, api_key=args.api_key)
    
    # Get the job description
    job_description = ""
    if args.job:
        try:
            with open(args.job, 'r') as file:
                job_description = file.read()
        except FileNotFoundError:
            print(f"Error: Job description file {args.job} not found.")
            return
    elif args.job_text:
        job_description = args.job_text
    else:
        print("Please provide a job description file (--job) or text (--job_text).")
        return
    
    # Generate the tailored resume
    tailored_resume = tailor.generate_tailored_resume(job_description)
    
    # Ensure output path has the correct extension
    output_path = args.output
    if not output_path.endswith(f'.{args.format}'):
        output_path = f"{output_path.rsplit('.', 1)[0] if '.' in output_path else output_path}.{args.format}"
    
    # Save the tailored resume
    tailor.save_resume(tailored_resume, output_path, args.format)

if __name__ == "__main__":
    main()
