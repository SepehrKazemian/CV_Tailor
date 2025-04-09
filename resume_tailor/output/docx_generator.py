"""
Word Document Generator Module

This module handles generating Word documents from tailored resumes.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_docx(resume_text: str, output_path: str) -> None:
    """
    Generate a Word document from a tailored resume.
    
    Args:
        resume_text: The tailored resume text
        output_path: Path to save the Word document to
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Split the resume into sections
    sections = resume_text.split("\n# ")
    
    # Process the first section (which doesn't start with #)
    if sections[0].startswith("# "):
        sections[0] = sections[0][2:]  # Remove the # prefix
    
    # Process each section
    for i, section in enumerate(sections):
        if i > 0:
            # For sections after the first one, the # was removed in the split
            section_parts = section.split("\n", 1)
            section_title = section_parts[0]
            section_content = section_parts[1] if len(section_parts) > 1 else ""
        else:
            # For the first section
            section_parts = section.split("\n", 1)
            section_title = section_parts[0]
            section_content = section_parts[1] if len(section_parts) > 1 else ""
        
        # Add section title
        title = doc.add_heading(section_title, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Add section content
        for paragraph_text in section_content.split("\n\n"):
            if paragraph_text.strip():
                # Check if this is a bullet point list
                if "•" in paragraph_text:
                    for bullet_point in paragraph_text.split("•"):
                        if bullet_point.strip():
                            p = doc.add_paragraph()
                            p.style = "List Bullet"
                            p.add_run(bullet_point.strip())
                else:
                    p = doc.add_paragraph()
                    p.add_run(paragraph_text.strip())
    
    # Save the document
    doc.save(output_path)
    print(f"Resume saved as Word document to {output_path}")


def estimate_pages(resume_text: str) -> int:
    """
    Estimate the number of pages a resume would take in a Word document.
    
    Args:
        resume_text: The resume text
        
    Returns:
        Estimated number of pages
    """
    # This is a rough estimate - actual page count depends on font, margins, etc.
    chars_per_page = 3500
    return max(1, (len(resume_text) + chars_per_page - 1) // chars_per_page)
